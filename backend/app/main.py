#!/usr/bin/env python3
"""PBL 导师工作台后端"""
import os, re, json, sqlite3, urllib.parse
from pathlib import Path
from fastapi import FastAPI, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

BASE = Path(__file__).resolve().parent.parent.parent
COURSES_DIR = BASE / "courses"
DB_PATH = BASE / "data" / "kb.db"

app = FastAPI(title="PBL 导师工作台", docs_url="/pbl-api/docs", openapi_url="/pbl-api/openapi.json")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

# ═══════════════ 课程 API ═══════════════

MODULES = [
    ("module1", "模块 1｜PBL 是什么", "认知建立：为什么是 PBL、探究式 vs 项目式、全景与思维"),
    ("module2", "模块 2｜设计项目", "核心模块：设计师思维、定义挑战、驱动性问题、以终为始、个性化、21 世纪技能、计划设计"),
    ("module3", "模块 3｜管理团队", "实施阶段：培养团队、避免危机、规划进程、构建文化"),
    ("module4", "模块 4｜评估与复盘", "评估成长、用好观众、复盘反思"),
    ("module5", "模块 5｜推进者进阶", "深度学习、思维转变、真实性、协作、创造力、失败原因"),
    ("module6", "模块 6｜线上 PBL", "差异化亮点：线上设计、7 原则、7 策略、协作与社群、评估分享"),
]

def parse_course_meta(path: Path):
    """从课件 md 提取元信息"""
    text = path.read_text(encoding="utf-8")
    title = ""
    m = re.search(r"^# (.+)$", text, re.M)
    if m:
        title = m.group(1).strip()
    # 学习目标（兼容 **学习目标**： 和 学习目标： 格式）
    objective = ""
    m = re.search(r"\*\*学习目标\*\*[：:](.+?)(?:\n|$)|学习目标[：:](.+?)(?:\n|$)", text)
    if m:
        objective = (m.group(1) or m.group(2) or "").strip()
    # 字数
    chars = len(re.sub(r"[\s#*`>|-]", "", text))
    return {"title": title, "objective": objective, "chars": chars}

@app.get("/pbl-api/courses/nav")
def course_nav():
    """返回全部课程的有序列表（供上一课/下一课导航）"""
    items = []
    for mod_id, _, _ in MODULES:
        mod_dir = COURSES_DIR / mod_id
        if mod_dir.exists():
            for f in sorted(mod_dir.glob("*.md")):
                meta = parse_course_meta(f)
                items.append({"id": f.stem, "title": meta["title"], "module": mod_id})
    return {"courses": items}

@app.get("/pbl-api/courses")
def list_courses():
    modules = []
    for mod_id, name, desc in MODULES:
        mod_dir = COURSES_DIR / mod_id
        courses = []
        if mod_dir.exists():
            for f in sorted(mod_dir.glob("*.md")):
                meta = parse_course_meta(f)
                course_id = f.stem
                courses.append({
                    "id": course_id,
                    "title": meta["title"],
                    "objective": meta["objective"],
                    "chars": meta["chars"],
                    "order": f.name.split("-")[0],
                })
        modules.append({"id": mod_id, "name": name, "desc": desc, "courses": courses})
    return {"modules": modules}

@app.get("/pbl-api/courses/{course_id}")
def get_course(course_id: str):
    for mod_id, _, _ in MODULES:
        p = COURSES_DIR / mod_id / f"{course_id}.md"
        if p.exists():
            content = p.read_text(encoding="utf-8")
            # 把「知识库检索"XXX"」转成可点击链接（跳转知识库页并自动搜索）
            import html as _html
            def _linkify(m):
                kw = m.group(1)
                href = f"/pbl/knowledge?q={urllib.parse.quote(kw)}"
                return f'<a class="kb-link" href="{href}" target="_blank">知识库检索「{_html.escape(kw)}」</a>'
            content = re.sub(r'知识库检索[“"]([^”"]+)[”"]', _linkify, content)
            return {"id": course_id, "module": mod_id, "content": content}
    return {"error": "not found"}

# ═══════════════ 知识库 API ═══════════════

def get_conn():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

@app.get("/pbl-api/kb/search")
def kb_search(q: str = Query(..., min_length=1), topic: str = "", limit: int = 20):
    conn = get_conn()
    q = q.strip()
    if not q:
        return {"results": []}
    params = []
    # trigram 索引：3 字及以上用 MATCH（精准+高亮），短词用 LIKE（覆盖）
    where_parts = []
    if len(q) >= 3:
        esc = re.sub(r'["*()]', ' ', q).strip()
        if esc:
            where_parts.append("books MATCH ?")
            params.append(f'"{esc}"')
    if not where_parts or len(q) < 3:
        where_parts.append("content LIKE ?")
        params.append(f"%{q}%")
    if topic:
        where_parts.append("topic = ?")
        params.append(topic)
    sql = "SELECT title, topic, page, snippet(books, -1, '[', ']', '…', 10) AS snip FROM books WHERE " + " AND ".join(where_parts) + " LIMIT ?"
    params.append(limit)
    try:
        rows = conn.execute(sql, params).fetchall()
        results = [{"title": r["title"], "topic": r["topic"], "page": r["page"], "snippet": r["snip"]} for r in rows]
    except Exception:
        # 最终兜底：纯 LIKE
        sql = "SELECT title, topic, page, substr(content, 1, 200) AS snip FROM books WHERE content LIKE ?"
        params2 = [f"%{q}%"]
        if topic:
            sql += " AND topic = ?"
            params2.append(topic)
        sql += " LIMIT ?"
        params2.append(limit)
        rows = conn.execute(sql, params2).fetchall()
        results = [{"title": r["title"], "topic": r["topic"], "page": r["page"], "snippet": r["snip"]} for r in rows]
    conn.close()
    return {"results": results, "total": len(results)}

@app.get("/pbl-api/kb/topics")
def kb_topics():
    conn = get_conn()
    rows = conn.execute("SELECT topic, COUNT(*) as cnt FROM books WHERE topic != '' GROUP BY topic ORDER BY cnt DESC").fetchall()
    conn.close()
    return {"topics": [{"name": r["topic"], "count": r["cnt"]} for r in rows]}

@app.get("/pbl-api/kb/books")
def kb_books():
    """9 本书列表（含页数、简介：取第 2 页内容为简介）"""
    conn = get_conn()
    rows = conn.execute("SELECT title, COUNT(*) as pages, MIN(page) as minp FROM books GROUP BY title ORDER BY pages DESC").fetchall()
    books = []
    for r in rows:
        # 找目录页或前几页做简介
        intro = ""
        pages = conn.execute("SELECT content FROM books WHERE title=? AND page < 15 AND LENGTH(content) > 60 ORDER BY page LIMIT 3", (r["title"],)).fetchall()
        if pages:
            intro = pages[-1]["content"].replace("\n", " ")[:120]
        books.append({"title": r["title"], "pages": r["pages"], "intro": intro})
    conn.close()
    return {"books": books}

@app.get("/pbl-api/kb/book")
def kb_book_page(title: str = Query(...), page: int = Query(1)):
    """按书名+页码取内容"""
    conn = get_conn()
    row = conn.execute("SELECT page, content FROM books WHERE title=? AND page=?", (title, page)).fetchone()
    if not row:
        # 找最近页
        row = conn.execute("SELECT page, content FROM books WHERE title=? ORDER BY ABS(page-?) LIMIT 1", (title, page)).fetchone()
    total = conn.execute("SELECT COUNT(*) FROM books WHERE title=?", (title,)).fetchone()[0]
    conn.close()
    return {"title": title, "page": row["page"], "total_pages": total, "content": row["content"] if row else ""}

# ═══════════════ 工具卡 API ═══════════════

TOOLCARDS_PATH = BASE / "data" / "toolcards.json"

def _load_toolcards():
    if TOOLCARDS_PATH.exists():
        import json as _json
        return _json.loads(TOOLCARDS_PATH.read_text(encoding="utf-8"))
    return []

@app.get("/pbl-api/tools")
def list_tools():
    """全部工具卡"""
    return {"tools": _load_toolcards()}

@app.get("/pbl-api/tools/{tool_id}")
def get_tool(tool_id: str):
    for t in _load_toolcards():
        if t["id"] == tool_id:
            return t
    return {"error": "not found"}

# ═══════════════ 项目库 API ═══════════════

BUILTIN_TEMPLATES = [
        {"id": "hz-museum", "title": "杭博暑期研学·文物会说话", "age": "8-12 岁", "duration": "3 小时",
         "driving_q": "如何为一件文物设计一场展览，让同龄人愿意来看？",
         "product": "迷你策展方案 + 讲解词", "scene": "杭州博物馆", "module": "模块 2 设计项目"},
        {"id": "liangzhu", "title": "良渚博物院·玉器密码", "age": "10-14 岁", "duration": "半天",
         "driving_q": "良渚的玉器为什么能'说话'？请你为一件玉器写一份身份档案。",
         "product": "玉器身份档案 + 讲解", "scene": "良渚博物院", "module": "模块 2 设计项目"},
        {"id": "westlake-water", "title": "西湖水质侦探", "age": "9-12 岁", "duration": "半天",
         "driving_q": "西湖的水质适合鱼类生存吗？我们如何验证并保护它？",
         "product": "水质报告 + 保护建议书", "scene": "西湖", "module": "模块 2 设计项目"},
        {"id": "budapest", "title": "布达佩斯展·跨时空对话", "age": "8-14 岁", "duration": "全天",
         "driving_q": "如果匈牙利文物会说话，它们最想告诉杭州孩子什么？",
         "product": "跨时空对话展览 + 导览", "scene": "杭州博物馆（布达佩斯展）", "module": "模块 6 线上 PBL"},
        {"id": "silk-road", "title": "丝绸之路·杭州出发", "age": "10-15 岁", "duration": "1 天",
         "driving_q": "丝绸如何改变世界？请你设计一条从杭州出发的'新丝绸之路'展线。",
         "product": "展线方案 + 沙盘", "scene": "中国丝绸博物馆", "module": "模块 2 设计项目"},
        {"id": "qiannian-song", "title": "南宋风雅·生活美学", "age": "8-12 岁", "duration": "半天",
         "driving_q": "南宋人怎样把日子过成诗？请你复刻一件南宋生活器物。",
         "product": "器物复刻 + 生活美学手册", "scene": "南宋官窑博物馆", "module": "模块 5 推进者进阶"},
    ]

def _init_template_db():
    conn = get_conn()
    conn.execute("CREATE TABLE IF NOT EXISTS custom_templates (id INTEGER PRIMARY KEY AUTOINCREMENT, title TEXT, age TEXT, duration TEXT, driving_q TEXT, product TEXT, scene TEXT, created_at TEXT DEFAULT CURRENT_TIMESTAMP)")
    conn.commit()
    conn.close()

@app.get("/pbl-api/templates")
def list_templates():
    """项目模板（内置 + 用户自存）"""
    _init_template_db()
    conn = get_conn()
    customs = conn.execute("SELECT * FROM custom_templates ORDER BY id DESC").fetchall()
    conn.close()
    custom_list = [{"id": f"c{t['id']}", "title": t["title"], "age": t["age"], "duration": t["duration"],
                    "driving_q": t["driving_q"], "product": t["product"], "scene": t["scene"],
                    "module": "导师自存", "custom": True} for t in customs]
    return {"templates": BUILTIN_TEMPLATES + custom_list}

class TemplateData(BaseModel):
    title: str
    age: str = ""
    duration: str = ""
    driving_q: str = ""
    product: str = ""
    scene: str = ""

@app.post("/pbl-api/templates/save")
def save_template(data: TemplateData):
    """把工坊产出存为项目模板"""
    _init_template_db()
    conn = get_conn()
    cur = conn.execute(
        "INSERT INTO custom_templates(title, age, duration, driving_q, product, scene) VALUES (?,?,?,?,?,?)",
        (data.title, data.age, data.duration, data.driving_q, data.product, data.scene))
    conn.commit()
    tid = cur.lastrowid
    conn.close()
    return {"ok": True, "id": f"c{tid}"}

# ═══════════════ 工坊 API ═══════════════

class WorkshopData(BaseModel):
    template_id: str = ""
    intent: str = ""
    driving_question: str = ""
    audience: str = ""
    age: str = ""
    duration: str = ""
    scene: str = ""
    product: str = ""
    evaluation: str = ""
    plan: str = ""
    objectives: str = ""
    evaluation_detail: str = ""
    resources: str = ""
    phases: list = []

class AIQuery(BaseModel):
    intent: str = ""
    audience: str = ""
    scene: str = ""
    count: int = 3

class AIAction(BaseModel):
    action: str  # intent|polish|elements|plan|questions|decompose|rubric
    intent: str = ""
    audience: str = ""
    scene: str = ""
    driving_question: str = ""
    age: str = ""
    duration: str = ""
    product: str = ""
    evaluation: str = ""
    plan: str = ""

PROMPTS = {
    "intent": """你是 PBL 研学课程设计专家。请根据以下信息，用一句话写出这门课的「意图声明」（不超过 60 字，说明希望学生带走什么能力）。
目标受众：{audience}
解决的真实问题：{scene}
要求：直接输出意图声明，不要解释。""",
    "polish": """你是 PBL 课程设计专家。请用 4 个标准（重要有意义/真实世界关联/开放有挑战/可持续探究）检验下面的驱动性问题，并给出改进后的版本。
当前驱动性问题：{driving_question}
课程受众：{audience}，场景：{scene}
输出格式：
问题诊断：<一句话指出最大问题>
改进版本：<改写后的驱动性问题>
只输出这两行。""",
    "elements": """你是 PBL 研学课程设计专家。请根据课程信息，给出立项五要素建议。
课程意图：{intent}
驱动性问题：{driving_question}
目标受众：{audience}
输出格式（每行一个，直接给出建议值）：
年龄：<建议年龄段>
时长：<建议时长>
场景：<建议场地>
成果产出：<建议的最终成果>
评估方式：<建议的评估方法>
只输出这五行。""",
    "decompose": """你是 PBL 课程设计专家。请把核心驱动问题分解成 3-4 个子问题，每个子问题对应一个任务和产出。
驱动性问题：{driving_question}
场景：{scene}
受众：{audience}
时长：{duration}
输出格式（严格 3-4 行，每行用 | 分隔 4 字段）：
子问题 | 核心任务（2-3步，用；分隔） | 学生产出 | 里程碑节点
示例：文物想说什么？| 听文物自述音频；分组认领文物角色卡；写一句"文物心声"独白 | 文物身份卡+独白稿 | 入项完成
必须贴合本课程。""",
    "rubric": """你是 PBL 评价设计专家。请把评价方案拆成结构化评价矩阵（4 列）。
课程场景：{scene}
受众：{audience}
评价方案原文：{evaluation}
输出格式（严格 3-4 行，每行用 | 分隔 4 字段）：
评价内容 | 证据 | 评价方式 | 评价时机
示例：策展内容准确性 | 文物身份卡 + 导览词 | 教师用评价表打分 + 同伴互评 | 第 2 阶段结束
必须贴合本课程。""",
    "detail": """你是 PBL 研学课程设计专家。请根据课程信息，补全课程详设（学习目标、评价方案、资源需求）。
课程意图：{intent}
驱动性问题：{driving_question}
目标受众：{audience}
时长：{duration}
场景：{scene}
成果：{product}
输出格式（每行一个部分，直接给出内容）：
学习目标：<知识目标；技能目标；素养目标>
评价方案：<形成性评价+终结性评价，怎么评何时评>
资源需求：<场地；物料；人员>
只输出这三行，每行以"学习目标：/评价方案：/资源需求："开头。""",
    "plan": """你是 PBL 研学课程设计专家。请为这门课设计 3 个实施阶段，每阶段一行，严格用 | 分隔 5 个字段，不要输出其他文字。
课程意图：{intent}
驱动性问题：{driving_question}
目标受众：{audience}
时长：{duration}
场景：{scene}
成果：{product}
每行格式：阶段名|时间|核心活动(2-3步，用；分隔)|学生产出|教师动作
示例：策展启蒙|1课时|入项游戏破冰；分组认领文物|文物认领卡|引导讨论发放入项材料
必须贴合课程主题。""",
}

def _call_deepseek(prompt):
    import urllib.request as _req, json as _json
    key = os.environ.get("DEEPSEEK_API_KEY", "")
    if not key:
        return None, "未配置 DEEPSEEK_API_KEY"
    body = _json.dumps({
        "model": "deepseek-chat",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.6,
        "max_tokens": 400,
        "stream": False
    }).encode()
    r = _req.Request("https://api.deepseek.com/chat/completions", data=body, headers={
        "Content-Type": "application/json",
        "Authorization": "Bearer " + key})
    try:
        resp = _json.loads(_req.urlopen(r, timeout=60).read())
        return resp["choices"][0]["message"]["content"].strip(), None
    except Exception as e:
        return None, str(e)

@app.post("/pbl-api/workshop/ai")
def ai_action(data: AIAction):
    """通用 AI 动作：intent|polish|elements|plan"""
    if data.action not in PROMPTS:
        return {"ok": False, "error": "未知 action"}
    prompt = PROMPTS[data.action].format(
        intent=data.intent or "（未填写）",
        audience=data.audience or "中小学生研学团",
        scene=data.scene or "博物馆研学",
        driving_question=data.driving_question or "（未填写）",
        duration=data.duration or "半天",
        product=data.product or "（未定）",
        evaluation=data.evaluation or "（未填写）")
    content, err = _call_deepseek(prompt)
    if err:
        return {"ok": False, "error": err}
    # 解析结果
    if data.action == "intent":
        return {"ok": True, "result": content}
    if data.action == "polish":
        return {"ok": True, "result": content}
    if data.action == "elements":
        lines = [l for l in content.splitlines() if l.strip()]
        result = {}
        for l in lines:
            if "：" in l:
                k, v = l.split("：", 1)
                k = k.strip().replace("年龄", "age").replace("时长", "duration").replace("场景", "scene").replace("成果产出", "product").replace("评估方式", "evaluation")
                if k in ("age", "duration", "scene", "product", "evaluation"):
                    result[k] = v.strip()
        return {"ok": True, "result": result}
    if data.action == "plan":
        # 解析 | 分隔的结构化阶段
        phases = []
        for line in content.splitlines():
            line = line.strip()
            if not line or line.startswith("阶段名") or "|" not in line:
                continue
            parts = [x.strip() for x in line.split("|")]
            phases.append({
                "name": parts[0],
                "time": parts[1] if len(parts) > 1 else "",
                "activities": parts[2] if len(parts) > 2 else "",
                "output": parts[3] if len(parts) > 3 else "",
                "teacher": parts[4] if len(parts) > 4 else "",
            })
        if phases:
            return {"ok": True, "result": phases}
        return {"ok": True, "result": content}
    if data.action == "detail":
        result = {}
        for l in content.splitlines():
            if "：" in l:
                k, v = l.split("：", 1)
                k = k.strip()
                if k == "学习目标": result["objectives"] = v.strip()
                elif k == "评价方案": result["evaluation_detail"] = v.strip()
                elif k == "资源需求": result["resources"] = v.strip()
        return {"ok": True, "result": result}
    if data.action == "decompose":
        phases = []
        for line in content.splitlines():
            line = line.strip()
            if not line or "|" not in line:
                continue
            parts = [x.strip() for x in line.split("|")]
            if len(parts) >= 3:
                phases.append({
                    "sub_q": parts[0],
                    "task": parts[1] if len(parts) > 1 else "",
                    "output": parts[2] if len(parts) > 2 else "",
                    "milestone": parts[3] if len(parts) > 3 else "",
                })
        return {"ok": True, "result": phases} if phases else {"ok": True, "result": content}
    if data.action == "rubric":
        rows = []
        for line in content.splitlines():
            line = line.strip()
            if not line or "|" not in line:
                continue
            parts = [x.strip() for x in line.split("|")]
            if len(parts) >= 3:
                rows.append({
                    "criterion": parts[0],
                    "evidence": parts[1] if len(parts) > 1 else "",
                    "method": parts[2] if len(parts) > 2 else "",
                    "timing": parts[3] if len(parts) > 3 else "",
                })
        return {"ok": True, "result": rows} if rows else {"ok": True, "result": content}
    return {"ok": True, "result": content}

@app.post("/pbl-api/workshop/ai-questions")
def ai_generate_questions(data: AIQuery):
    """用 DeepSeek 生成候选驱动性问题"""
    import urllib.request as _urlopen_req
    import json as _json
    key = os.environ.get("DEEPSEEK_API_KEY", "")
    if not key:
        return {"ok": False, "error": "未配置 DEEPSEEK_API_KEY"}
    prompt = f"""你是 PBL（项目式学习）课程设计专家。请根据以下信息设计 {data.count} 个高质量的驱动性问题。

课程意图：{data.intent or "（未填写）"}
目标受众：{data.audience or "中小学生研学团"}
场景：{data.scene or "博物馆/城市研学"}（示例参考：杭州博物馆、西湖、良渚博物院、布达佩斯展）

要求：
1. 每个问题都满足：重要有意义、真实世界关联、开放有挑战、可持续探究
2. 不要用"什么"类教科书式问题；不要是"是/否"就能回答的
3. 尽量本地化、具体化，适合研学场景
4. 问题要能让学生"用"知识而不是"背"知识

输出格式（严格按此格式，每个问题一行）：
Q1: <问题>
Q2: <问题>
Q3: <问题>
只输出 Q1/Q2/Q3 三行，不要其他内容。"""
    body = _json.dumps({
        "model": "deepseek-chat",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.7,
        "max_tokens": 500,
        "stream": False
    }).encode()
    req = _urlopen_req.Request("https://api.deepseek.com/chat/completions", data=body, headers={
        "Content-Type": "application/json",
        "Authorization": "Bearer " + key
    })
    try:
        resp = _json.loads(_urlopen_req.urlopen(req, timeout=60).read())
        content = resp["choices"][0]["message"]["content"]
        questions = [l.strip()[4:].strip() for l in content.splitlines() if l.strip().startswith("Q")]
        if not questions:
            questions = [l.strip() for l in content.splitlines() if l.strip()][:data.count]
        return {"ok": True, "questions": questions[:data.count]}
    except Exception as e:
        return {"ok": False, "error": str(e)}

@app.post("/pbl-api/workshop/export")
def workshop_export(data: WorkshopData):
    """生成课程方案 .docx"""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    # 标题
    title = doc.add_heading("PBL 研学课程方案", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("浸思研学 · PBL 课程开发工坊出品")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    def section(name, content):
        doc.add_heading(name, level=1)
        doc.add_paragraph(content or "（待补充）")

    section("一、课程意图（为什么做这门课）", data.intent)
    section("二、驱动性问题（项目的心脏）", data.driving_question)
    doc.add_heading("三、立项五要素", level=1)
    items = [
        ("目标受众", data.audience),
        ("适用年龄", data.age),
        ("时长", data.duration),
        ("场景", data.scene),
        ("成果产出", data.product),
        ("评估方式", data.evaluation),
    ]
    for k, v in items:
        p = doc.add_paragraph()
        r = p.add_run(f"{k}：")
        r.bold = True
        p.add_run(v or "（待补充）")
    section("四、学习目标", data.objectives)
    section("五、评价方案", data.evaluation_detail)
    section("六、资源需求", data.resources)
    doc.add_heading("七、实施计划", level=1)
    if data.phases:
        from docx.shared import Cm as _Cm
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(["阶段", "时间", "核心活动", "学生产出", "教师动作"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True if hdr[i].paragraphs[0].runs else None
        for idx, ph in enumerate(data.phases, 1):
            row = table.add_row().cells
            row[0].text = ph.get("name") or f"阶段{idx}"
            row[1].text = ph.get("time") or ""
            row[2].text = ph.get("activities") or ""
            row[3].text = ph.get("output") or ""
            row[4].text = ph.get("teacher") or ""
        for row in table.rows:
            row.cells[0].width = _Cm(2.5)
            row.cells[1].width = _Cm(2)
            row.cells[2].width = _Cm(5.5)
            row.cells[3].width = _Cm(3)
            row.cells[4].width = _Cm(3)
    else:
        doc.add_paragraph(data.plan or "（待补充）")

    # 保存到内存返回
    import io
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    from fastapi.responses import Response
    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="pbl_course_plan.docx"'}
    )

# ═══════════════ 静态前端 ═══════════════

FRONT_DIST = BASE / "frontend" / "dist"
if FRONT_DIST.exists():
    app.mount("/pbl", StaticFiles(directory=FRONT_DIST, html=True), name="pbl-front")

# ═══════════════ 统计 & 反馈 ═══════════════

def _init_meta_db():
    conn = get_conn()
    conn.execute("CREATE TABLE IF NOT EXISTS page_views (id INTEGER PRIMARY KEY AUTOINCREMENT, page TEXT, course_id TEXT DEFAULT '', ts TEXT DEFAULT CURRENT_TIMESTAMP)")
    conn.execute("CREATE TABLE IF NOT EXISTS feedback (id INTEGER PRIMARY KEY AUTOINCREMENT, course_id TEXT, rating INTEGER, note TEXT DEFAULT '', ts TEXT DEFAULT CURRENT_TIMESTAMP)")
    conn.commit()
    conn.close()

@app.post("/pbl-api/stats/track")
def track_view(data: dict):
    _init_meta_db()
    conn = get_conn()
    conn.execute("INSERT INTO page_views(page, course_id) VALUES (?,?)", (data.get("page", ""), data.get("course_id", "")))
    conn.commit()
    conn.close()
    return {"ok": True}

@app.get("/pbl-api/stats")
def get_stats():
    _init_meta_db()
    conn = get_conn()
    total_views = conn.execute("SELECT COUNT(*) FROM page_views").fetchone()[0]
    unique_pages = conn.execute("SELECT COUNT(DISTINCT page) FROM page_views").fetchone()[0]
    today = conn.execute("SELECT COUNT(*) FROM page_views WHERE date(ts) = date('now','localtime')").fetchone()[0]
    # 热门课程
    hot = conn.execute("SELECT course_id, COUNT(*) c FROM page_views WHERE course_id != '' GROUP BY course_id ORDER BY c DESC LIMIT 5").fetchall()
    # 反馈统计
    fb_total = conn.execute("SELECT COUNT(*) FROM feedback").fetchone()[0]
    fb_pos = conn.execute("SELECT COUNT(*) FROM feedback WHERE rating >= 4").fetchone()[0]
    conn.close()
    return {
        "views": total_views, "pages": unique_pages, "today": today,
        "hot_courses": [{"id": r["course_id"], "count": r["c"]} for r in hot],
        "feedback": {"total": fb_total, "positive": fb_pos}
    }

@app.post("/pbl-api/feedback")
def submit_feedback(data: dict):
    _init_meta_db()
    conn = get_conn()
    conn.execute("INSERT INTO feedback(course_id, rating, note) VALUES (?,?,?)",
                 (data.get("course_id", ""), data.get("rating", 0), data.get("note", "")))
    conn.commit()
    conn.close()
    return {"ok": True}

@app.get("/pbl-api/health")
def health():
    return {"status": "ok", "courses_dir": str(COURSES_DIR), "kb": str(DB_PATH)}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8100)
